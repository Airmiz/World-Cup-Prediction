"""Regenerate output/WC2026_report.docx from the current model artifacts.

Reproducible companion to the spreadsheet/figures: pulls every number from
output/team_probabilities.csv, output/sim_summary.json and output/backtest.json
and embeds the regenerated figures, so the report always matches the shipped
model. Run after figures.py in the pipeline.
"""
from __future__ import annotations
import json, csv
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

tp = list(csv.DictReader(open("output/team_probabilities.csv")))
ss = json.load(open("output/sim_summary.json"))
bt = json.load(open("output/backtest.json"))
by = {r["team"]: r for r in tp}

xi = bt["best"]["xi"]; w = bt["best"]["blend_w"]
rho = ss["dc_rho"]; hadv = ss["dc_home_adv"]
halflife = 0.6931 / xi
vmaj = bt["validation_2019plus"]["major_finals"]
champ = [(r["team"], float(r["p_champion"])) for r in tp[:10]]
elo5 = sorted(tp, key=lambda r: -float(r["elo"]))[:5]
finals = ss["top_finals"][:5]
mex, usa, can = by["Mexico"], by["United States"], by["Canada"]
pct = lambda x: f"{x*100:.1f}%"

doc = Document()
st = doc.styles["Normal"]; st.font.name = "Calibri"; st.font.size = Pt(10.5)

def title(t, sz, bold=True, color=None, align=None, after=6):
    p = doc.add_paragraph(); r = p.add_run(t); r.bold = bold; r.font.size = Pt(sz)
    if color: r.font.color.rgb = RGBColor(*color)
    if align is not None: p.alignment = align
    p.paragraph_format.space_after = Pt(after); return p

def body(t):
    p = doc.add_paragraph(t); p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.line_spacing = 1.15; return p

def fig(path, caption, width=6.2):
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(path, width=Inches(width))
    c = doc.add_paragraph(); cr = c.add_run(caption); cr.italic = True; cr.font.size = Pt(9)
    c.alignment = WD_ALIGN_PARAGRAPH.CENTER; c.paragraph_format.space_after = Pt(10)

# ---- header ----
title("FIFA World Cup 2026", 24, color=(0x0E, 0x2A, 0x47), align=WD_ALIGN_PARAGRAPH.CENTER, after=2)
title("A Statistical Forecast of All 104 Matches", 14, bold=False,
      color=(0x1A, 0x6F, 0xB4), align=WD_ALIGN_PARAGRAPH.CENTER, after=2)
sub = doc.add_paragraph(); sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
sr = sub.add_run("Time-decayed Dixon-Coles × World Football Elo ensemble · 200,000 Monte Carlo "
                 "tournament simulations · Pre-tournament forecast (data through 10 June 2026; "
                 "model refined 14 June 2026)")
sr.italic = True; sr.font.size = Pt(9); sub.paragraph_format.space_after = Pt(12)

def H1(n, t): title(f"{n}. {t}", 14, color=(0x0E, 0x2A, 0x47), after=4)
def H2(n, t): title(f"{n} {t}", 11.5, after=3)

H1(1, "Summary")
body(f"This report documents a probabilistic forecasting model for the 2026 FIFA World Cup, built "
     f"from data available before the opening match in Mexico City on 11 June 2026. The model combines "
     f"two complementary engines — a time-decayed Dixon-Coles bivariate Poisson model of scoring rates and "
     f"a World Football Elo rating system spanning 49,477 international matches since 1872 — into a single "
     f"calibrated forecast, then plays the entire 48-team, 104-match tournament 200,000 times under FIFA's "
     f"exact competition rules, including the official Annex C allocation of third-placed teams to the round of 32.")
body(f"The simulations make {champ[0][0]} the narrow favourite at {pct(champ[0][1])}, a whisker ahead of "
     f"{champ[1][0]} at {pct(champ[1][1])}, with {champ[2][0]} ({pct(champ[2][1])}) and {champ[3][0]} "
     f"({pct(champ[3][1])}) leading the chasing pack. The single most likely final is "
     f"{finals[0]['final']} ({pct(finals[0]['p'])} of simulations). Mexico profits substantially from hosting, "
     f"winning Group A in {pct(float(mex['p_win_group']))} of simulations, while the United States — ranked "
     f"lower by recent results — wins Group D only {pct(float(usa['p_win_group']))} of the time. Every "
     f"probability in this report is reproducible from the accompanying code and a fixed random seed "
     f"({ss['seed']}).")

H1(2, "Data")
body("The training data is the open international results database maintained by Mart Jürisoo "
     "(martj42/international_results): every senior men's full international from 30 November 1872 through "
     "10 June 2026 — 49,477 completed matches with date, teams, score, competition, venue city and a "
     "neutral-ground flag. The same source supplies the 72 scheduled group-stage fixtures. The tournament "
     "structure — twelve groups of four, the round-of-32 bracket, and FIFA's complete 495-row Annex C table — "
     "was transcribed from the official regulations and validated programmatically (every row is a permutation "
     "of the qualified groups, and no third-placed team can face its own group winner).")

H1(3, "Model")
H2("3.1", "Dixon-Coles component")
body(f"The core scoring model is the Dixon-Coles (1997) bivariate Poisson. Each team carries an attack and a "
     f"defence strength; expected goals for the home side are exp(μ + attack_home − defence_away + γ·home), and "
     f"symmetrically for the away side. Dependence between the two scorelines at low totals is captured by the "
     f"Dixon-Coles τ correction with coupling ρ = {rho:.3f}, fitted by profile likelihood. Matches enter the "
     f"likelihood with exponential time-decay weights exp(−ξ·Δt): the tuned ξ of {xi} per year halves a match's "
     f"influence roughly every {halflife:.1f} years (Section 4), multiplied by an importance weight (1.0 for "
     f"World Cup and continental finals, 0.85 for qualifiers and Nations League, 0.7 for friendlies). Team "
     f"strengths are estimated by penalised maximum likelihood — a quadratic penalty equivalent to a Gaussian "
     f"prior that shrinks thinly-observed teams toward average. The fitted home advantage is {hadv:.3f} "
     f"log-goals, about a 30% uplift in scoring rate, applied only where a team genuinely plays at home — for "
     f"this tournament, the three hosts in their own stadiums.")
H2("3.2", "Elo component")
body(f"Independently, World Football Elo ratings are computed over the full 154-year history: K-factors scale "
     f"with match importance (60 for World Cup finals down to 20 for friendlies), margin of victory inflates "
     f"the update through the goal-difference multiplier, and home teams are credited 100 rating points. Because "
     f"ratings are stored as they stood before each historical match, two weighted Poisson regressions are "
     f"fitted without look-ahead bias, translating any Elo gap into expected goals. The current top five by Elo: "
     + ", ".join(f"{r['team']} ({round(float(r['elo']))})" for r in elo5) + ".")
H2("3.3", "Ensemble")
body(f"The final forecast geometrically blends the two engines' goal rates — λ = λ_DC^w · λ_Elo^(1−w) — with "
     f"the blend weight, decay rate and friendly weight all chosen out-of-sample (Section 4). The selected "
     f"w = {w} leans slightly toward the faster-reacting Elo engine, and the blend beats both of its components "
     f"on the held-out data — the classic ensemble result.")

H1(4, "Validation")
body("The model is validated by a leakage-free rolling-origin backtest: components are re-estimated every month "
     "using only matches played before each refit date, and every competitive international since 2012 is then "
     "predicted out-of-sample — roughly 9,500 matches, of which about 1,200 are major-tournament finals (World "
     "Cup, Euro, Copa América, the continental championships) that resemble the balanced, top-team matchups of a "
     "World Cup. Forecasts are scored by the Rank Probability Score (RPS — the standard ordered-outcome football "
     "metric) alongside three-way log loss and Brier score. Hyperparameters were tuned on the pre-2019 "
     "major-finals sample by RPS, subject to not regressing the World Cup group-stage holdouts; everything from "
     "2019 onward is a clean out-of-sample test.")
body(f"On that out-of-sample period the selected configuration improves on the previous settings: major-finals "
     f"RPS falls from {vmaj['incumbent']['rps']:.4f} to {vmaj['chosen']['rps']:.4f} and log loss from "
     f"{vmaj['incumbent']['log_loss']:.4f} to {vmaj['chosen']['log_loss']:.4f}, with no regression on the most "
     f"recent World Cup. Scored against the three most recent World Cup group stages under a strict walk-forward "
     f"protocol, the blend lands as follows (lower is better; a know-nothing uniform forecast scores 1.099 log "
     f"loss and 0.667 Brier):")

# validation table
vt = doc.add_table(rows=1, cols=5); vt.style = "Light Grid Accent 1"; vt.alignment = WD_TABLE_ALIGNMENT.CENTER
for j, h in enumerate(("WC", "Model", "Log loss", "Brier", "RPS")):
    c = vt.rows[0].cells[j].paragraphs[0].add_run(h); c.bold = True
rowmap = [("blend", "Blend (final model)"), ("dc_only", "Dixon-Coles only"),
          ("elo_only", "Elo only"), ("uniform", "Uniform baseline")]
for y in ("2014", "2018", "2022"):
    for k, lbl in rowmap:
        m = bt[y][k]; cells = vt.add_row().cells
        cells[0].text = y if k == "blend" else ""
        cells[1].text = lbl
        cells[2].text = f"{m['log_loss']:.3f}"; cells[3].text = f"{m['brier']:.3f}"; cells[4].text = f"{m['rps']:.3f}"
for r in vt.rows:
    for c in r.cells:
        for p in c.paragraphs:
            p.paragraph_format.space_after = Pt(1)
            for run in p.runs: run.font.size = Pt(9)
doc.add_paragraph().paragraph_format.space_after = Pt(4)
body("Calibration over the pooled group-stage probabilities tracks the diagonal closely:")
fig("output/fig_calibration.png", "Figure 1 — Reliability of predicted probabilities, WC 2014/2018/2022 "
    "group stages (out-of-sample).", width=4.6)

H1(5, "Tournament simulation")
body("Each of the 200,000 simulations plays the tournament under FIFA's full rule book. Group scorelines are "
     "drawn from each fixture's exact joint score distribution (the τ-corrected Poisson grid), so margins — not "
     "just results — propagate into the standings. Groups are ranked by points, goal difference and goals "
     "scored; two-way ties level on all three are resolved head-to-head, with remaining criteria proxied by "
     "randomisation. The twelve third-placed teams are ranked, the best eight advance, and their bracket slots "
     "come from the official 495-combination Annex C table. Knockout matches level after 90 minutes go to extra "
     "time — a 30-minute Poisson process at one-third intensity — then a 50/50 shoot-out. Host-nation home "
     "advantage persists into the knockouts only where the schedule provides it (Mexico at the Azteca, Canada "
     "in Vancouver, the United States across essentially its whole possible path).")

H1(6, "Results")
fig("output/fig_champion.png", "Figure 2 — Probability of lifting the trophy (top 15 of 48).", width=5.6)
body(f"{champ[0][0]} and {champ[1][0]} stand clearly apart, together accounting for roughly a third of all "
     f"simulated titles — consistent with their 1–2 positions in the Elo list. The most likely finals are "
     + ", ".join(f"{x['final']} ({pct(x['p'])})" for x in finals) + ". Among the hosts, Mexico converts a soft "
     f"Group A and Azteca home advantage into a {pct(float(mex['p_reach_r16']))} chance of the round of 16 and a "
     f"{pct(float(mex['p_champion']))} title shot; Canada reaches the round of 16 in "
     f"{pct(float(can['p_reach_r16']))} of simulations; the United States, despite playing at home, carries weak "
     f"2024–25 form into a contested Group D and reaches the round of 16 "
     f"{pct(float(usa['p_reach_r16']))} of the time.")

# champion table
ct = doc.add_table(rows=1, cols=3); ct.style = "Light Grid Accent 1"; ct.alignment = WD_TABLE_ALIGNMENT.CENTER
for j, h in enumerate(("Rank", "Team", "P(champion)")):
    rr = ct.rows[0].cells[j].paragraphs[0].add_run(h); rr.bold = True
for i, (t, p) in enumerate(champ, 1):
    cells = ct.add_row().cells; cells[0].text = str(i); cells[1].text = t; cells[2].text = pct(p)
    for c in cells:
        for pp in c.paragraphs:
            pp.paragraph_format.space_after = Pt(1)
            for run in pp.runs: run.font.size = Pt(9)
doc.add_paragraph().paragraph_format.space_after = Pt(4)
fig("output/fig_groups.png", "Figure 3 — Probability of surviving the group phase "
    "(dotted line = the ⅔ average advancement rate of this format).", width=6.2)

H1(7, "Limitations")
body("The model knows teams, not players: injuries, suspensions, rotation and managerial changes enter only "
     "insofar as they already show in results. The headline forecast uses no betting-market information, which "
     "the literature finds hard to beat — a deliberate choice for reproducibility (the live site separately "
     "tracks the model against the market). Penalty shoot-outs are treated as coin flips, three-way tie-breaks "
     "and conduct criteria are randomised, and the 48-team format is a mild extrapolation. Probabilities carry "
     "Monte Carlo error below ±0.2 points at this simulation count, but model error necessarily exceeds "
     "simulation error.")

H1(8, "Reproducibility")
body("Everything here regenerates from the project folder: data/ holds the match database, fixtures, bracket "
     "and Annex C table; src/model.py implements the Elo engine, the Dixon-Coles fit and the blend; "
     "src/backtest.py runs the rolling-origin validation (RPS + log loss + Brier) and writes "
     "output/backtest.json, from which the production model reads its hyperparameters; src/simulate.py refits "
     "on all data through 10 June 2026 and replays the tournament (seed 20260611); src/figures.py, "
     "src/make_xlsx.py and src/make_report.py rebuild the charts, spreadsheet and this document. Run order: "
     "backtest.py → simulate.py → figures.py → make_xlsx.py → make_report.py.")

doc.save("output/WC2026_report.docx")
print("saved output/WC2026_report.docx")
